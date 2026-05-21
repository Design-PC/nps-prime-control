from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
LOGO = ROOT / "public" / "brand" / "prime-control-logo.png"
DOCX_PATH = OUT_DIR / "Prime_Control_NPS_Documentacao_Executiva_Tecnica.docx"

BLUE = "003F7D"
DARK = "102A43"
MUTED = "5C6F82"
LIGHT_BLUE = "E9F2FB"
LIGHT_GRAY = "F4F7FB"
LINE = "D7E2EC"
ORANGE = "F58220"
WHITE = "FFFFFF"
GREEN = "067647"
RED = "B42318"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = DARK, size: float = 9.4) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = LINE, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=120, start=180, bottom=120, end=180) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def set_run_font(run, size=None, color=DARK, bold=None, italic=None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, bold=False, color=DARK, size=10.5, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=11.5, color=DARK, bold=True)
    return p


def add_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=10.2, color=DARK)
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=10.2, color=DARK)
    return p


def add_callout(doc, title: str, body: str, fill: str = LIGHT_BLUE, accent: str = BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=accent, size="8")
    set_cell_margins(table, top=140, bottom=140, start=180, end=180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=DARK)
    add_para(doc, "", after=4)
    return table


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, title in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, BLUE)
        set_cell_text(cell, title, bold=True, color=WHITE, size=9.2)
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9.1)
            if widths:
                cells[i].width = Inches(widths[i])
    add_para(doc, "", after=4)
    return table


def add_metric_strip(doc, metrics: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=len(metrics))
    table.autofit = False
    set_table_borders(table, color="BFD2E4")
    set_cell_margins(table, top=120, bottom=120, start=120, end=120)
    for i, (label, value) in enumerate(metrics):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT_BLUE if i % 2 == 0 else "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        set_run_font(r, size=8.4, color=MUTED, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=11, color=BLUE, bold=True)
    add_para(doc, "", after=4)
    return table


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.34)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, DARK, 7, 3),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.86))
    header_table.autofit = False
    set_table_borders(header_table, color=WHITE, size="0")
    set_cell_margins(header_table, top=0, bottom=0, start=0, end=0)
    left = header_table.cell(0, 0)
    right = header_table.cell(0, 1)
    left.width = Inches(2.5)
    right.width = Inches(4.36)
    if LOGO.exists():
        p_logo = left.paragraphs[0]
        p_logo.paragraph_format.space_after = Pt(0)
        p_logo.add_run().add_picture(str(LOGO), width=Inches(1.5))
    p_title = right.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_title.paragraph_format.space_after = Pt(0)
    r = p_title.add_run("Plataforma NPS Corporativa")
    set_run_font(r, size=8.8, color=BLUE, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = footer_p.add_run("Prime Control | Projeto NPS Corporativo | Confidencial | Página ")
    set_run_font(r1, size=8.5, color=MUTED)
    add_page_number(footer_p)


def add_cover(doc: Document) -> None:
    cover_table = doc.add_table(rows=1, cols=2)
    cover_table.autofit = False
    set_table_borders(cover_table, color=WHITE, size="0")
    set_cell_margins(cover_table, top=0, bottom=60, start=0, end=0)
    logo_cell = cover_table.cell(0, 0)
    title_cell = cover_table.cell(0, 1)
    logo_cell.width = Inches(2.65)
    title_cell.width = Inches(4.21)
    if LOGO.exists():
        p_logo = logo_cell.paragraphs[0]
        p_logo.paragraph_format.space_after = Pt(0)
        p_logo.add_run().add_picture(str(LOGO), width=Inches(2.05))
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_title.paragraph_format.space_after = Pt(0)
    r_title = p_title.add_run("Plataforma NPS Corporativa")
    set_run_font(r_title, size=12, color=BLUE, bold=True)

    add_para(doc, "DOCUMENTAÇÃO EXECUTIVA E TÉCNICA", color=BLUE, bold=True, size=11, before=44, after=8)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Plataforma de Pesquisa NPS Corporativa")
    set_run_font(r, size=28, color=DARK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Prime Control | Clientes Ativos B2B")
    set_run_font(r, size=16, color=BLUE, bold=True)

    add_callout(
        doc,
        "Objetivo do projeto",
        "Transformar a pesquisa NPS em uma experiência digital moderna, rastreável e orientada a conversão, preservando integralmente as perguntas aprovadas e criando inteligência comportamental para Marketing, CS, liderança e dados.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    add_metric_strip(
        doc,
        [
            ("Foco", "Adesão"),
            ("Experiência", "Multi-step"),
            ("Dados", "Dashboard + planilha"),
            ("Analytics", "Clarity + eventos"),
        ]
    )

    add_para(doc, f"Versão: 1.2 | Atualizado em: {date.today().strftime('%d/%m/%Y')}", size=10, color=MUTED, before=18, after=34)
    add_para(doc, "Documento consolidado a partir da auditoria, estratégia, PRD técnico, roadmap, backlog, operação de campanha e implementação inicial do MVP.", size=9.6, color=MUTED)
    doc.add_page_break()


def add_summary(doc: Document) -> None:
    add_heading(doc, "1. Sumário executivo", 1)
    add_callout(
        doc,
        "Recomendação principal",
        "Construir uma plataforma própria de NPS, em vez de depender do Microsoft Forms ou Typeform como solução final. A pesquisa precisa ser tratada como produto estratégico de experiência do cliente, com dados comportamentais, dashboard executivo e integração com CRM.",
        fill="FFF2E8",
        accent=ORANGE,
    )
    add_para(
        doc,
        "A pesquisa atual possui perguntas metodologicamente aprovadas, mas a experiência visual e operacional reduz adesão. O problema central não é a nota, e sim a baixa visibilidade da jornada: hoje a empresa sabe quem respondeu ou não, mas não sabe quem abriu, clicou, iniciou, abandonou, em qual etapa houve fricção ou quais clientes silenciosos exigem follow-up.",
    )
    add_para(
        doc,
        "A solução proposta reorganiza a experiência em etapas curtas, preserva todas as perguntas, melhora percepção de esforço, captura comportamento do respondente e gera inteligência acionável para Marketing, CS e liderança.",
    )
    add_metric_strip(
        doc,
        [
            ("Entrada", "Abertura e clique"),
            ("Conversão", "Início e conclusão"),
            ("Fricção", "Drop-off e hesitação"),
            ("Ação", "CS e liderança"),
        ]
    )
    add_heading(doc, "Decisões estratégicas já tomadas", 2)
    for item in [
        "Preservar integralmente as perguntas existentes, sem reescrita ou enviesamento.",
        "Usar experiência multi-step para reduzir sobrecarga cognitiva.",
        "Salvar progresso e permitir retomada por token individual.",
        "Gerar planilha executiva e dashboard operacional com respostas e comportamento.",
        "Integrar Microsoft Clarity para mapas de calor e replay de sessão.",
        "Preparar integração com HubSpot para envio, segmentação e follow-up.",
    ]:
        add_bullet(doc, item)


def add_context_history(doc: Document) -> None:
    add_heading(doc, "2. Histórico e contexto do projeto", 1)
    add_para(
        doc,
        "A pesquisa NPS já existia em ciclos anteriores, incluindo uma versão utilizada em 2025. Para 2026, a liderança revisou a abordagem e definiu um novo conjunto de perguntas em formato de planilha/Excel. O print usado como referência neste projeto representa esse material de definição das perguntas, não a interface final desejada para o cliente.",
    )
    add_callout(
        doc,
        "Premissa importante",
        "O trabalho atual não parte da ideia de mudar as perguntas. O objetivo é transformar perguntas já definidas em uma experiência digital mais clara, rastreável e estratégica, com maior adesão e melhor leitura de comportamento.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )
    rows = [
        ["2025", "Pesquisa anterior", "Baixa visibilidade: resposta ou não resposta, sem leitura de jornada."],
        ["2026", "Perguntas revisadas", "Perguntas definidas em planilha; interface final ainda precisa ser desenhada."],
        ["Projeto atual", "Plataforma NPS", "Experiência multi-step, analytics, dashboard, planilha e integração com CRM."],
    ]
    add_table(doc, ["Momento", "Contexto", "Implicação"], rows, widths=[1.1, 1.75, 3.65])


def add_toc(doc: Document) -> None:
    add_heading(doc, "3. Estrutura do documento", 1)
    rows = [
        ["1", "Sumário executivo", "Decisão, racional e visão de negócio."],
        ["2", "Histórico e contexto", "Versão 2025, revisão 2026 e premissas."],
        ["3", "Auditoria da pesquisa atual", "Riscos de UX, fricção e baixa adesão."],
        ["4", "Redesenho da experiência", "Arquitetura multi-step e racional psicológico."],
        ["5", "Jornada comportamental", "Riscos, objeções e mitigação por etapa."],
        ["6", "Analytics e dados", "Clarity, eventos, planilha, dashboard e tracking."],
        ["7", "PRD técnico", "Frontend, backend, banco, segurança e integrações."],
        ["8", "Roadmap e backlog", "Fases de implementação e prioridades."],
        ["9", "Operação da campanha", "Processo, papéis, rituais e métricas de sucesso."],
        ["10", "MVP atual e governança", "O que já foi feito, próximos passos e atualização contínua."],
    ]
    add_table(doc, ["#", "Seção", "Finalidade"], rows, widths=[0.45, 2.35, 3.7])


def add_audit(doc: Document) -> None:
    add_heading(doc, "4. Auditoria estratégica da pesquisa atual", 1)
    add_para(doc, "O formulário atual se apresenta como uma planilha corporativa. Isso comunica esforço, burocracia e baixa modernidade, especialmente para clientes ativos B2B e respondentes executivos ocupados.")
    rows = [
        ["Exposição total", "Todas as perguntas aparecem de uma vez.", "Aumenta percepção de esforço e adia o início."],
        ["Layout de tabela", "Escalas em colunas estreitas e bordas fortes.", "Dificulta mobile, aumenta erro de clique e reduz percepção premium."],
        ["Sem progressão", "Não há etapa atual, barra de progresso ou tempo restante.", "Eleva incerteza e abandono nos blocos finais."],
        ["Identificação inicial", "Dados manuais aparecem antes do valor percebido.", "Gera fricção e desperdiça dados já disponíveis no CRM."],
        ["Campos abertos secos", "Campos aparecem sem apoio ou microcopy neutra.", "Reduz qualidade das respostas e pode gerar bloqueio."],
        ["Analytics insuficiente", "Mede apenas respondeu ou não respondeu.", "Impede otimização e identificação de fricções reais."],
    ]
    add_table(doc, ["Problema", "Evidência", "Impacto na adesão"], rows, widths=[1.45, 2.25, 2.8])
    add_heading(doc, "Pontos prováveis de abandono", 2)
    for item in [
        "Antes de iniciar, ao perceber o tamanho total do formulário.",
        "Na identificação, caso o usuário precise preencher dados já conhecidos.",
        "No primeiro campo aberto após a nota NPS.",
        "Em blocos repetitivos de escala, especialmente no mobile.",
        "Na última seção, por concentração de perguntas abertas.",
    ]:
        add_bullet(doc, item)


def add_experience(doc: Document) -> None:
    add_heading(doc, "5. Redesenho da experiência", 1)
    add_callout(
        doc,
        "Princípio de UX",
        "A pesquisa deve parecer menor do que é. Não reduzimos perguntas; reduzimos percepção de esforço por meio de etapas, progresso, salvamento e linguagem executiva.",
    )
    rows = [
        ["0", "Landing", "Propósito, tempo estimado, confiança e CTA."],
        ["1", "Identificação", "Dados pré-preenchidos via HubSpot/token."],
        ["2", "Relacionamento e Satisfação", "NPS e motivo da nota."],
        ["3", "Percepção de Valor", "Avaliação de valor, relevância e resultados."],
        ["4", "Qualidade Operacional", "Entregas, prazos, clareza, atendimento e resposta."],
        ["5", "Inovação, Transformação e Futuro", "Última etapa com visão estratégica e campos abertos."],
        ["6", "Conclusão", "Confirmação executiva e encerramento."],
    ]
    add_table(doc, ["Etapa", "Nome", "Papel na jornada"], rows, widths=[0.55, 2.25, 3.7])
    add_heading(doc, "Racional psicológico", 2)
    for item in [
        "Chunking: blocos menores reduzem sobrecarga cognitiva.",
        "Goal gradient effect: progresso visível aumenta motivação para concluir.",
        "Perceived control: autosave e retomada reduzem ansiedade.",
        "Executive respect: tempo estimado e linguagem objetiva sinalizam respeito pelo tempo.",
        "Neutralidade metodológica: microcopy apoia o preenchimento sem sugerir nota.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Decisão sobre a ordem da escala", 2)
    add_para(
        doc,
        "A escala deve ser apresentada visualmente de forma crescente, da esquerda para a direita. Para NPS, o padrão mais reconhecido é 0 a 10; para perguntas de satisfação, quando a metodologia aprovada usa 1 a 10, a apresentação recomendada é 1 a 10.",
    )
    rows = [
        ["1 → 10 ou 0 → 10", "Fluxo mental natural de leitura", "Reduz erro de clique e transmite neutralidade."],
        ["10 → 1 ou 10 → 0", "Ordem invertida", "Aumenta carga cognitiva, pode parecer indução à nota alta e prejudica mobile."],
        ["Sem destaque prévio do 10", "Neutralidade visual", "Evita sinalizar que a empresa espera uma resposta positiva."],
    ]
    add_table(doc, ["Formato", "Leitura do usuário", "Efeito esperado"], rows, widths=[1.55, 2.2, 2.75])


def add_journey(doc: Document) -> None:
    add_heading(doc, "6. Jornada comportamental do respondente", 1)
    rows = [
        ["Recebe e-mail", "Ignorar", "Assunto claro, remetente confiável e personalização."],
        ["Abre e-mail", "Leitura superficial", "Texto curto, tempo estimado e propósito."],
        ["Clica no link", "Medo de formulário longo", "Landing leve, sem expor todas as perguntas."],
        ["Inicia", "Adiar", "CTA claro e promessa de 3 a 5 minutos."],
        ["Confirma dados", "Fricção", "Pré-preenchimento via HubSpot."],
        ["Responde NPS", "Dúvida de escala", "Escala crescente e rótulos claros."],
        ["Campos abertos", "Bloqueio textual", "Microcopy neutra: uma frase já ajuda."],
        ["Finaliza", "Fadiga tardia", "Mensagem de última etapa e progresso visível."],
    ]
    add_table(doc, ["Momento", "Risco", "Mitigação"], rows, widths=[1.65, 1.9, 2.95])
    add_heading(doc, "Objeções humanas esperadas", 2)
    rows = [
        ["Não tenho tempo agora", "Tempo estimado, autosave e retomada."],
        ["Isso vai virar cobrança?", "Texto claro sobre uso responsável e melhoria contínua."],
        ["Não sei o que escrever", "Aceitar respostas curtas e microcopy neutra."],
        ["Já dei a nota, por que responder o resto?", "Separar blocos por propósito e mostrar progresso."],
    ]
    add_table(doc, ["Objeção", "Resposta da experiência"], rows, widths=[2.25, 4.25])


def add_analytics(doc: Document) -> None:
    add_heading(doc, "7. Analytics, planilha e dashboard", 1)
    add_para(doc, "A nova plataforma deve medir entrada, comportamento, abandono, conversão e qualidade da resposta. A camada de dados é parte central do produto, não um complemento.")
    rows = [
        ["Microsoft Clarity", "Mapas de calor, replay de sessão, rage clicks, dead clicks, comportamento mobile.", "Entender como o usuário se comporta visualmente."],
        ["PostHog", "Eventos, funis, coortes, tempo por etapa, abandono por pergunta.", "Entender onde o funil quebra."],
        ["GA4", "Origem, dispositivo, navegador, UTM e tráfego.", "Entender de onde o usuário veio."],
        ["Banco próprio", "Respostas, sessões, tokens, progresso e eventos.", "Garantir confiabilidade e exportação."],
        ["Dashboard", "Participação, abandono, NPS, clientes silenciosos e risco.", "Apoiar liderança e CS."],
        ["Planilha", "CSV/XLSX de respostas e comportamento.", "Permitir análise, auditoria e compartilhamento."],
    ]
    add_table(doc, ["Camada", "O que mede", "Por que importa"], rows, widths=[1.4, 2.75, 2.35])
    add_heading(doc, "Explicação em linguagem executiva", 2)
    rows = [
        ["Clarity", "Mostra a tela como o cliente usou: onde clicou, onde travou e onde desistiu.", "É como assistir a experiência real, sem depender de suposição."],
        ["PostHog", "Conta os eventos da jornada: iniciou, avançou, abandonou, retomou e concluiu.", "Mostra exatamente em qual etapa a conversão quebra."],
        ["GA4", "Mostra origem e contexto do acesso: e-mail, dispositivo, navegador e campanha.", "Ajuda a entender se o problema está no tráfego ou no formulário."],
        ["Dashboard", "Transforma dados em visão executiva: participação, abandono, NPS e contas de risco.", "Facilita decisão e follow-up."],
    ]
    add_table(doc, ["Ferramenta", "Em termos simples", "Valor para liderança"], rows, widths=[1.25, 3.0, 2.25])
    add_heading(doc, "Taxonomia de eventos prioritários", 2)
    rows = [
        ["nps_landing_viewed", "Carregamento da landing", "campaign_id, recipient_id, device, source"],
        ["nps_survey_started", "Clique em iniciar", "time_to_start, contact_role"],
        ["nps_step_viewed", "Etapa exibida", "step_id, step_name, progress_percent"],
        ["nps_question_answered", "Resposta registrada", "question_id, category, time_to_answer"],
        ["nps_step_completed", "Avanço de etapa", "step_id, time_on_step, answered_count"],
        ["nps_survey_abandoned", "Saída/inatividade", "last_step, last_question, elapsed_time"],
        ["nps_survey_resumed", "Retorno pelo token", "last_step, days_since_last_activity"],
        ["nps_survey_completed", "Envio final", "total_time, nps_score, completion_path"],
    ]
    add_table(doc, ["Evento", "Trigger", "Properties"], rows, widths=[2.15, 1.85, 2.5])
    add_heading(doc, "Planilha de respostas", 2)
    add_para(doc, "A exportação CSV/XLSX deve conter respostas e sinais comportamentais, não apenas nota final.")
    for item in [
        "Identificação: empresa, contato, cargo, área e token/campaign_id.",
        "Status: convidado, abriu, clicou, iniciou, abandonou, retomou ou concluiu.",
        "Respostas: nota NPS, escalas por bloco e campos abertos.",
        "Comportamento: etapa abandonada, tempo por etapa, tempo total e dispositivo.",
        "Qualidade: campos abertos preenchidos, tamanho de comentário e resposta parcial.",
    ]:
        add_bullet(doc, item)


def add_dashboards(doc: Document) -> None:
    add_heading(doc, "8. Dashboards executivos", 1)
    rows = [
        ["Saúde da campanha", "Enviados, abertos, cliques, iniciados, concluídos, completion rate.", "Liderança e Marketing"],
        ["Drop-off e atrito", "Abandono por etapa, tempo por pergunta, friction score.", "Produto, UX e Dados"],
        ["Clientes silenciosos", "Não abriu, abriu e não clicou, iniciou e abandonou.", "CS e Gestão de Contas"],
        ["Inteligência NPS", "NPS por segmento, detratores, neutros, promotores e temas.", "Liderança e CS"],
        ["Qualidade de resposta", "Campos abertos, comentários úteis, respostas muito curtas.", "Marketing, Produto e Dados"],
    ]
    add_table(doc, ["Dashboard", "Principais visões", "Usuário interno"], rows, widths=[1.55, 3.45, 1.5])
    add_callout(
        doc,
        "Score de fricção",
        "Combina abandono, tempo acima do percentil 75, rage clicks, dead clicks, revisitas e campos ignorados para priorizar melhorias de UX.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )


def add_technical_prd(doc: Document) -> None:
    add_heading(doc, "9. PRD técnico", 1)
    add_heading(doc, "Arquitetura recomendada", 2)
    rows = [
        ["Frontend", "Next.js, TypeScript, Tailwind, Shadcn UI", "Experiência multi-step responsiva."],
        ["Banco", "Supabase ou PostgreSQL", "Persistência de respostas, sessões e eventos."],
        ["CRM", "HubSpot", "Envio, segmentação, status e lembretes."],
        ["Analytics", "Clarity, PostHog, GA4", "Comportamento visual, funil e origem."],
        ["Dashboard", "Power BI, Metabase ou dashboard interno", "Leitura executiva e operacional."],
        ["Segurança", "Token seguro, HTTPS, LGPD", "Controle de acesso e minimização de dados."],
    ]
    add_table(doc, ["Camada", "Tecnologia", "Responsabilidade"], rows, widths=[1.35, 2.3, 2.85])
    add_heading(doc, "Rotas e componentes", 2)
    for item in [
        "/nps/[token]: landing personalizada.",
        "/nps/[token]/survey: wizard da pesquisa.",
        "/nps/[token]/complete: confirmação de conclusão.",
        "Componentes: SurveyLayout, ProgressBar, RatingScale, OpenTextQuestion, AutosaveIndicator, ResumeBanner e CompletionState.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Modelo de dados", 2)
    rows = [
        ["nps_campaigns", "Campanha, período, status e versão."],
        ["nps_recipients", "Contato, empresa, token hash e status da jornada."],
        ["nps_sessions", "Sessão, etapa atual, progresso e tempo total."],
        ["nps_questions", "Perguntas aprovadas, tipo, categoria e ordem."],
        ["nps_answers", "Respostas numéricas/textuais, tempo e revisitas."],
        ["nps_events", "Eventos comportamentais e properties."],
    ]
    add_table(doc, ["Tabela", "Finalidade"], rows, widths=[2.0, 4.5])
    add_heading(doc, "LGPD e segurança", 2)
    for item in [
        "Token aleatório, único e com expiração; armazenar hash do token, não token puro.",
        "HTTPS obrigatório e acesso restrito ao dashboard.",
        "Mascaramento de dados sensíveis no Microsoft Clarity.",
        "Separação entre análise agregada e visualização individual autorizada.",
        "Política clara sobre uso das respostas e finalidade de melhoria contínua.",
    ]:
        add_bullet(doc, item)


def add_roadmap(doc: Document) -> None:
    add_heading(doc, "10. Roadmap, backlog e plano de ação", 1)
    rows = [
        ["Fase 0", "Alinhamento", "Congelar perguntas, regras metodológicas, campos HubSpot e permissões."],
        ["Fase 1", "MVP de alta conversão", "Landing, wizard, token, autosave, retomada e persistência."],
        ["Fase 2", "Analytics comportamental", "Clarity, PostHog, eventos, heatmaps e session replay."],
        ["Fase 3", "Dashboard e planilha", "Funil, abandono, clientes silenciosos e exportação CSV/XLSX."],
        ["Fase 4", "Integração HubSpot", "Status no CRM, segmentos e lembretes comportamentais."],
        ["Fase 5", "Otimização CRO", "Testes A/B de landing, CTA e agrupamento, sem influenciar nota."],
        ["Fase 6", "Inteligência preditiva", "Score de risco, temas qualitativos e alertas para CS."],
    ]
    add_table(doc, ["Fase", "Nome", "Entregas"], rows, widths=[0.85, 1.75, 3.9])
    add_heading(doc, "Backlog P0", 2)
    for item in [
        "Criar banco e APIs de sessão/respostas.",
        "Persistir respostas reais e progresso.",
        "Gerar exportação CSV/XLSX.",
        "Criar dashboard interno inicial.",
        "Integrar Microsoft Clarity com mascaramento de dados.",
        "Preparar tokens individuais para envio via HubSpot.",
    ]:
        add_bullet(doc, item)


def add_operation(doc: Document) -> None:
    add_heading(doc, "11. Processo operacional da campanha", 1)
    rows = [
        ["Product Owner", "Escopo, perguntas, priorização e critérios de sucesso."],
        ["Marketing/CRM", "HubSpot, e-mails, listas, UTMs e lembretes."],
        ["Customer Success", "Contas estratégicas, follow-up e clientes silenciosos."],
        ["Engenharia", "Plataforma, segurança, integrações e disponibilidade."],
        ["Dados/BI", "Eventos, dashboards, qualidade e leitura executiva."],
    ]
    add_table(doc, ["Papel", "Responsabilidade"], rows, widths=[2.0, 4.5])
    add_heading(doc, "Cadência recomendada", 2)
    for item in [
        "Antes da campanha: validar base HubSpot, tokens, tracking, dashboard e QA mobile.",
        "Durante a campanha: daily de 15 minutos para funil, abandono, contas estratégicas e bloqueios.",
        "Após a campanha: relatório executivo, análise de temas, clientes de risco e backlog de otimização.",
    ]:
        add_bullet(doc, item)


def add_mvp_status(doc: Document) -> None:
    add_heading(doc, "12. Status do MVP atual", 1)
    add_para(doc, "A primeira versão navegável já foi criada para validar a experiência do respondente e agora já possui uma camada funcional de dados para MVP local.")
    rows = [
        ["Implementado", "Next.js, landing, fluxo multi-step, perguntas preservadas, autosave, APIs internas, persistência local, dashboard MVP, exportação CSV, logo e identidade visual."],
        ["Ainda pendente", "Banco de produção, autenticação do admin, importação de respondentes, XLSX, HubSpot, PostHog, GA4 e publicação no subdomínio."],
        ["Decisão técnica", "Evoluir o MVP para plataforma própria, usando Typeform apenas como benchmark ou fallback emergencial."],
    ]
    add_table(doc, ["Status", "Descrição"], rows, widths=[1.65, 4.85])
    add_heading(doc, "URLs de referência no ambiente local", 2)
    for item in [
        "Home técnica do protótipo: http://localhost:3000",
        "Landing do respondente: http://localhost:3000/nps/demo-prime-control",
        "Pesquisa multi-step: http://localhost:3000/nps/demo-prime-control/survey",
        "Dashboard interno: http://localhost:3000/admin",
        "Exportação CSV: http://localhost:3000/api/admin/export.csv",
    ]:
        add_bullet(doc, item)


def add_current_implementation(doc: Document) -> None:
    add_heading(doc, "13. Implementação atual: dados, dashboard e Clarity", 1)
    add_para(
        doc,
        "A plataforma deixou de ser apenas um protótipo visual e passou a salvar respostas em uma camada de dados local para MVP. Essa camada será substituída por Supabase/PostgreSQL em produção, mantendo a mesma lógica de APIs.",
    )
    rows = [
        ["APIs internas", "Sessão, início, salvamento de respostas, conclusão e eventos."],
        ["Persistência MVP", "Arquivo local JSON para acelerar validação funcional."],
        ["Dashboard interno", "Visão de convidados, iniciados, concluídos, completion rate e abandono por etapa."],
        ["Exportação CSV", "Arquivo com identificação, status, datas, etapa e respostas."],
        ["Clarity", "Script preparado via NEXT_PUBLIC_CLARITY_PROJECT_ID."],
    ]
    add_table(doc, ["Entrega", "Descrição"], rows, widths=[1.65, 4.85])
    add_heading(doc, "Como ativar o Clarity", 2)
    for item in [
        "Criar ou abrir o projeto no Microsoft Clarity para nps.primecontrol.com.br.",
        "Copiar o Project ID do projeto.",
        "Criar o arquivo .env.local com NEXT_PUBLIC_CLARITY_PROJECT_ID=<id_do_projeto>.",
        "Reiniciar o servidor para o script entrar em vigor.",
        "Validar no Clarity se as sessões, heatmaps e replays começaram a aparecer.",
    ]:
        add_bullet(doc, item)


def add_governance(doc: Document) -> None:
    add_heading(doc, "14. Governança e atualização contínua", 1)
    add_para(
        doc,
        "Este documento deve acompanhar a evolução do projeto. Sempre que uma decisão de produto, UX, dados, tecnologia ou operação for tomada, a documentação executiva deve ser atualizada para manter uma fonte única de entendimento.",
    )
    rows = [
        ["Mudança de escopo", "Atualizar objetivo, impacto, fase do roadmap e critérios de aceite."],
        ["Nova decisão de UX", "Registrar racional, impacto esperado na adesão e cuidados contra viés."],
        ["Nova integração", "Registrar ferramenta, finalidade, dados trafegados e responsável."],
        ["Mudança na pesquisa", "Registrar origem da mudança, versão e validação metodológica."],
        ["Resultado de campanha", "Adicionar leitura executiva, aprendizados, fricções e próximos ajustes."],
    ]
    add_table(doc, ["Quando atualizar", "O que registrar"], rows, widths=[2.0, 4.5])
    add_callout(
        doc,
        "Próximo marco documental",
        "Após a implementação da camada de dados, o documento deve incluir prints do dashboard, exemplo de planilha exportada, eventos reais capturados e leitura inicial do Clarity.",
        fill="FFF2E8",
        accent=ORANGE,
    )


def build() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)
    core = doc.core_properties
    core.title = "Plataforma de Pesquisa NPS Corporativa Prime Control"
    core.subject = "Documentação executiva e técnica"
    core.author = "Prime Control"
    core.comments = "Documento consolidado para apresentação executiva e consulta técnica."

    add_cover(doc)
    add_summary(doc)
    add_context_history(doc)
    add_toc(doc)
    add_audit(doc)
    doc.add_page_break()
    add_experience(doc)
    add_journey(doc)
    doc.add_page_break()
    add_analytics(doc)
    add_dashboards(doc)
    doc.add_page_break()
    add_technical_prd(doc)
    add_roadmap(doc)
    doc.add_page_break()
    add_operation(doc)
    add_mvp_status(doc)
    add_current_implementation(doc)
    add_governance(doc)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
